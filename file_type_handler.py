"""
File Type Handler Module
=======================

This module provides specialized handlers for different file types (PDF, Word, Excel).
Each handler implements a common interface for metadata operations.

Key Programming Concepts Used:
- Classes: Base class and derived classes for file type handlers
- Inheritance: Type-specific handlers inherit from base class
- Polymorphism: Common interface for all file types
- Exception Handling: Try-except blocks for error management
- Type Checking: Validate file types and metadata
"""

# Standard library imports
import os
import re
import datetime
import traceback

# Third-party imports for different file formats
from PIL import Image, ExifTags
from PIL.ExifTags import TAGS, GPSTAGS
import piexif

# PDF handling libraries
try:
    import PyPDF2
    # For backward compatibility with different PyPDF2 versions
    if hasattr(PyPDF2, 'PdfReader'):
        from PyPDF2 import PdfReader, PdfWriter
    else:
        from PyPDF2 import PdfFileReader as PdfReader, PdfFileWriter as PdfWriter
    PDF_SUPPORT = True
except ImportError:
    print("Warning: PyPDF2 library not found. PDF support will be limited.")
    PDF_SUPPORT = False

# Word document handling
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    print("Warning: python-docx library not found. Word document support will be limited.")
    DOCX_SUPPORT = False

# Excel file handling
try:
    import openpyxl
    EXCEL_SUPPORT = True
except ImportError:
    print("Warning: openpyxl library not found. Excel spreadsheet support will be limited.")
    EXCEL_SUPPORT = False

# CLASS Definition
class BaseFileHandler:
    """
    Base class for all file type handlers.
    
    This class defines the common interface for all file type handlers
    and provides default implementations where possible.
    """
    
    def __init__(self):
        """Initialize BaseFileHandler instance."""
        self.file_extension = []
        self.display_name = "Generic File"
    
    def validate_file(self, file_path):
        """
        Validate if file exists and has the correct extension.
        
        Args:
            file_path (str): Path to the file to validate
            
        Returns:
            bool: True if file is valid, False otherwise
            
        Raises:
            ValueError: If the file is corrupted or invalid
        """
        # Check if file exists
        if not os.path.exists(file_path):
            return False
            
        # Check if it has the correct extension
        if not any(file_path.lower().endswith(ext) for ext in self.file_extension):
            return False
            
        return True
    
    def extract_metadata(self, file_path):
        """
        Extract metadata from file.
        
        Args:
            file_path (str): Path to file
            
        Returns:
            dict: Dictionary of metadata
            
        Raises:
            NotImplementedError: Must be implemented by subclasses
        """
        raise NotImplementedError("Subclasses must implement extract_metadata()")
    
    def remove_metadata(self, file_path, metadata_type):
        """
        Remove specified metadata from file.
        
        Args:
            file_path (str): Path to file
            metadata_type (str): Type of metadata to remove
            
        Returns:
            tuple: (modified file object, list of removed tags, new metadata dict)
            
        Raises:
            NotImplementedError: Must be implemented by subclasses
        """
        raise NotImplementedError("Subclasses must implement remove_metadata()")
    
    def add_metadata(self, file_path, metadata_dict):
        """
        Add or modify metadata in file.
        
        Args:
            file_path (str): Path to file
            metadata_dict (dict): Metadata to add/modify
            
        Returns:
            tuple: (modified file object, list of added/modified tags, new metadata dict)
            
        Raises:
            NotImplementedError: Must be implemented by subclasses
        """
        raise NotImplementedError("Subclasses must implement add_metadata()")
    
    def save_modified_file(self, original_path, file_obj, exif_dict, count):
        """
        Save modified file with clean metadata.
        
        Args:
            original_path (str): Path to original file
            file_obj: Modified file object
            exif_dict: Additional metadata dictionary (if applicable)
            count (int): File counter for unique filename
            
        Returns:
            str: Path to newly saved file
            
        Raises:
            NotImplementedError: Must be implemented by subclasses
        """
        raise NotImplementedError("Subclasses must implement save_modified_file()")
    
    def get_metadata_menu_options(self):
        """
        Get menu options specific to this file type.
        
        Returns:
            list: List of (option_name, metadata_type) tuples
        """
        # Default options
        return [
            ("Remove All Metadata", "all")
        ]


# CLASS Definition
class JPGHandler(BaseFileHandler):
    """Handler for JPG/JPEG image files."""
    
    def __init__(self):
        """Initialize JPGHandler instance."""
        super().__init__()
        self.file_extension = ['.jpg', '.jpeg']
        self.display_name = "JPEG Image"
    
    def validate_file(self, file_path):
        """
        Validate if file is a valid JPG.
        
        Args:
            file_path (str): Path to the file to validate
            
        Returns:
            bool: True if file is valid, False otherwise
            
        Raises:
            ValueError: If the file is corrupted or invalid
        """
        if not super().validate_file(file_path):
            return False
            
        # Attempt to open the file as an image - simple validation
        try:
            with Image.open(file_path) as img:
                # Try to access basic image properties to verify it's a valid image
                img.size
                return True
        except Exception as e:
            print(f"Warning: File validation issue: {e}")
            print("We'll attempt to process the file anyway.")
            return True
    
    def extract_metadata(self, file_path):
        """
        Extract metadata from JPG file with improved error handling and fallbacks.
        
        Args:
            file_path (str): Path to file
            
        Returns:
            dict: Dictionary of metadata
        """
        if not os.path.exists(file_path):
            return {"Error": "File not found"}

        try:
            # Enhanced metadata extraction with multiple methods
            readable_exif = {}
            
            with Image.open(file_path) as img:
                # Method 1: Try PIL's _getexif method
                exif = None
                try:
                    exif = img._getexif()
                    if exif:
                        for tag_id, value in exif.items():
                            try:
                                tag = TAGS.get(tag_id, str(tag_id))
                                
                                # Handle byte strings
                                if isinstance(value, bytes):
                                    try:
                                        value = value.decode('utf-8', 'ignore')
                                    except:
                                        value = f"[Binary Data: {len(value)} bytes]"
                                        
                                readable_exif[tag] = value
                            except Exception as e:
                                continue  # Skip problematic tags
                except Exception as e:
                    print(f"Warning: Error using PIL _getexif method: {e}")
                
                # Method 2: Try piexif directly if first method didn't find metadata
                if not readable_exif and 'exif' in img.info:
                    try:
                        exif_dict = piexif.load(img.info['exif'])
                        
                        # Convert piexif format to a readable dict
                        for ifd_name in ('0th', '1st', 'Exif', 'GPS', 'Interop'):
                            if ifd_name in exif_dict:
                                for tag_id, value in exif_dict[ifd_name].items():
                                    try:
                                        # Get tag name based on IFD
                                        if ifd_name == 'GPS' and tag_id in GPSTAGS:
                                            tag = f"GPS_{GPSTAGS[tag_id]}"
                                        elif tag_id in TAGS:
                                            tag = TAGS[tag_id]
                                        else:
                                            tag = f"{ifd_name}_{tag_id}"
                                            
                                        # Handle byte strings
                                        if isinstance(value, bytes):
                                            try:
                                                value = value.decode('utf-8', 'ignore')
                                            except:
                                                value = f"[Binary Data: {len(value)} bytes]"
                                                
                                        readable_exif[tag] = value
                                    except Exception as e:
                                        continue  # Skip problematic tags
                    except Exception as e:
                        print(f"Warning: Error using piexif method: {e}")
                
                # Method 3: Try basic image properties if no metadata found
                if not readable_exif:
                    # Get basic image properties
                    readable_exif["Format"] = img.format
                    readable_exif["Mode"] = img.mode
                    readable_exif["Size"] = f"{img.width} x {img.height}"
                    
                    # Check file properties
                    file_stats = os.stat(file_path)
                    mtime = datetime.datetime.fromtimestamp(file_stats.st_mtime)
                    readable_exif["File Modified"] = mtime.strftime("%Y-%m-%d %H:%M:%S")
                    readable_exif["File Size"] = f"{file_stats.st_size} bytes"
                
                return readable_exif
        except Exception as e:
            print(f"Warning: Error extracting metadata from JPG: {e}")
            # Return basic file info instead of empty dict
            try:
                file_stats = os.stat(file_path)
                mtime = datetime.datetime.fromtimestamp(file_stats.st_mtime)
                return {
                    "Info": "Could not extract image metadata",
                    "File Path": file_path,
                    "File Size": f"{file_stats.st_size} bytes",
                    "File Modified": mtime.strftime("%Y-%m-%d %H:%M:%S")
                }
            except:
                return {"Info": "Could not extract image metadata or file information"}
    
    def remove_metadata(self, file_path, metadata_type):
        """
        Remove specified metadata from JPG with improved reliability and effectiveness.
        
        Args:
            file_path (str): Path to file
            metadata_type (str): Type of metadata to remove
            
        Returns:
            tuple: (image object, list of removed tags, exif dictionary)
        """
        try:
            # Open image
            img = Image.open(file_path)
            removed_tags = []
            
            # Check if the image has EXIF data to remove
            has_exif = False
            if hasattr(img, '_getexif') and img._getexif():
                has_exif = True
            elif 'exif' in img.info:
                has_exif = True
            
            # Create a new image without any metadata
            # This is the most reliable way to remove all metadata
            clean_img = Image.new(img.mode, img.size)
            clean_img.paste(img)
            
            # Create an empty EXIF dictionary for compatibility with save_modified_file
            exif_dict = {
                '0th': {}, '1st': {}, 'Exif': {}, 'GPS': {}, 'Interop': {}
            }
            
            # Record what we've removed for logging purposes
            if has_exif:
                if metadata_type in ["datetime", "all"]:
                    removed_tags.append("DateTime Information")
                
                if metadata_type in ["gps", "all"]:
                    removed_tags.append("GPS Information")
                    
                if metadata_type == "all":
                    removed_tags.append("ALL_METADATA")
                
                # If no specific tags were identified but we know there was EXIF data
                if not removed_tags:
                    removed_tags.append("EXIF Data")
            else:
                # No EXIF data found to remove
                removed_tags.append("No EXIF data found to remove")
            
            return clean_img, removed_tags, exif_dict
            
        except Exception as e:
            # If any error occurs, try to continue with a clean image
            print(f"Warning: Error during metadata removal: {e}")
            try:
                # Create a new clean image
                img = Image.open(file_path)
                clean_img = Image.new(img.mode, img.size)
                clean_img.paste(img)
                exif_dict = {'0th': {}, '1st': {}, 'Exif': {}, 'GPS': {}, 'Interop': {}}
                return clean_img, ["ALL_METADATA"], exif_dict
            except Exception as e2:
                # If that fails too, raise the error
                raise Exception(f"Error processing image file: {e2}")
    
    def add_metadata(self, file_path, metadata_dict):
        """
        Add or modify metadata in JPG file.
        
        Args:
            file_path (str): Path to file
            metadata_dict (dict): Metadata to add/modify 
                                 {'tag_group': {'tag_name': value}}
            
        Returns:
            tuple: (image object, list of added/modified tags, exif dictionary)
        """
        try:
            # Open image
            img = Image.open(file_path)
            
            # Initialize EXIF dictionary - empty by default for reliability
            exif_dict = {
                '0th': {}, '1st': {}, 'Exif': {}, 'GPS': {}, 'Interop': {}
            }
            
            # Try to load existing EXIF data
            try:
                if 'exif' in img.info:
                    loaded_exif = piexif.load(img.info['exif'])
                    # Copy each IFD dictionary
                    for ifd in exif_dict:
                        if ifd in loaded_exif:
                            exif_dict[ifd] = loaded_exif[ifd]
            except Exception as e:
                print(f"Warning: Could not load original EXIF: {e}")
                # Continue with empty exif_dict
            
            modified_tags = []
            
            # Process each metadata item to add/modify
            for group_name, tags in metadata_dict.items():
                group = None
                
                # Map group name to piexif group
                if group_name.lower() == 'image':
                    group = '0th'
                elif group_name.lower() == 'exif':
                    group = 'Exif'
                elif group_name.lower() == 'gps':
                    group = 'GPS'
                else:
                    # Skip invalid groups
                    continue
                
                for tag_name, value in tags.items():
                    # Find tag ID
                    tag_id = None
                    
                    # Look up tag ID based on group and name
                    if group == '0th':
                        for tid, tname in TAGS.items():
                            if tname == tag_name:
                                tag_id = tid
                                break
                    elif group == 'Exif':
                        for tid, tname in TAGS.items():
                            if tname == tag_name:
                                tag_id = tid
                                break
                    elif group == 'GPS':
                        for tid, tname in GPSTAGS.items():
                            if tname == tag_name:
                                tag_id = tid
                                break
                    
                    if tag_id is not None:
                        # Handle different value types
                        if isinstance(value, str):
                            # Convert string to bytes for EXIF
                            encoded_value = value.encode('utf-8')
                        else:
                            # Use value as is
                            encoded_value = value
                        
                        # Add/modify tag
                        exif_dict[group][tag_id] = encoded_value
                        modified_tags.append(f"{group}_{tag_name}")
            
            return img, modified_tags, exif_dict
            
        except Exception as e:
            print(f"Warning: Error during metadata addition: {e}")
            try:
                # Open image again and return with original state
                img = Image.open(file_path)
                exif_dict = {'0th': {}, '1st': {}, 'Exif': {}, 'GPS': {}, 'Interop': {}}
                return img, [], exif_dict
            except Exception as e2:
                raise Exception(f"Error opening image file: {e2}")
    
    def save_modified_file(self, original_path, img, exif_dict, count):
        """
        Save JPG with modified metadata, ensuring metadata is properly removed if requested.
        
        Args:
            original_path (str): Path to original image
            img: Image object
            exif_dict: EXIF dictionary
            count (int): File counter for unique filename
            
        Returns:
            str: Path to newly saved file
        """
        try:
            # Create new filename
            base_name = os.path.splitext(original_path)[0]
            new_file = f"{base_name}_modified_{count}.jpg"
            
            # Check if the new file path is too long
            max_length = 250  # Adjust as needed for your OS
            if len(new_file) > max_length:
                # Truncate the base name to fit within limits
                allowed_base_length = max_length - len(f"_modified_{count}.jpg") - 10
                truncated_base = os.path.splitext(original_path)[0][:allowed_base_length]
                new_file = f"{truncated_base}_modified_{count}.jpg"
            
            # Check if this is a metadata removal operation
            is_empty_exif = all(len(exif_dict[ifd]) == 0 for ifd in exif_dict)
            
            if is_empty_exif:
                # Save without any EXIF data
                img.save(new_file, "JPEG", quality=100)
            else:
                # Save with EXIF data
                try:
                    exif_bytes = piexif.dump(exif_dict)
                    img.save(new_file, "JPEG", exif=exif_bytes, quality=100)
                except Exception as e:
                    # If EXIF dumping fails, save without EXIF
                    print(f"Warning: Could not save EXIF data: {e}")
                    img.save(new_file, "JPEG", quality=100)
            
            return new_file
            
        except Exception as e:
            raise Exception(f"Error saving modified image: {e}")
    
    def get_metadata_menu_options(self):
        """
        Get menu options specific to JPG files.
        
        Returns:
            list: List of (option_name, metadata_type) tuples
        """
        return [
            ("Remove Date/Time Information", "datetime"),
            ("Remove GPS Location", "gps"),
            ("Remove All Metadata", "all"),
            ("Add/Modify Date/Time", "mod_datetime"),
            ("Add/Modify GPS Location", "mod_gps"),
            ("Add/Modify Camera Information", "mod_camera")
        ]


class PDFHandler(BaseFileHandler):
    """Handler for PDF files."""
    
    def __init__(self):
        """Initialize PDFHandler instance."""
        super().__init__()
        self.file_extension = ['.pdf']
        self.display_name = "PDF Document"
    
    def validate_file(self, file_path):
        """
        Validate if file is a valid PDF.
        
        Args:
            file_path (str): Path to the file to validate
            
        Returns:
            bool: True if file is valid, False otherwise
            
        Raises:
            ValueError: If the file is corrupted or invalid
        """
        if not super().validate_file(file_path):
            return False
            
        # Basic validation - just check for PDF header
        try:
            with open(file_path, 'rb') as f:
                # Check PDF header
                header = f.read(5)
                if header != b'%PDF-':
                    print("Warning: File doesn't have a valid PDF header.")
                    return False
            return True
        except Exception as e:
            print(f"Warning: Error validating PDF file: {e}")
            return False
    
    def extract_metadata(self, file_path):
        """
        Extract metadata from PDF file with improved compatibility.
        
        Args:
            file_path (str): Path to file
            
        Returns:
            dict: Dictionary of metadata
        """
        # Check if PDF support is available
        if not PDF_SUPPORT:
            return {"Error": "PyPDF2 library not available for PDF metadata extraction"}
            
        try:
            # Try to open and read the PDF
            with open(file_path, 'rb') as f:
                # Create PDF reader object
                reader = PdfReader(f)
                
                metadata = {}
                
                # Get document info - handling newer PyPDF2 versions (3.0.0+)
                if hasattr(reader, 'metadata') and reader.metadata:
                    # For newer PyPDF2 versions
                    info = reader.metadata
                    
                    # Get all available attributes
                    for key in [
                        'author', 'creator', 'producer', 'subject', 
                        'title', 'creation_date', 'modification_date'
                    ]:
                        try:
                            if hasattr(info, key):
                                value = getattr(info, key)
                                if value:
                                    metadata[key] = str(value)
                        except Exception as e:
                            print(f"Warning: Error reading PDF metadata field {key}: {e}")
                
                # If no metadata found with the newer method, try the older method
                if not metadata and hasattr(reader, 'getDocumentInfo'):
                    # For older PyPDF2 versions
                    info = reader.getDocumentInfo()
                    if info:
                        for key, value in info.items():
                            # Strip the leading '/' from keys
                            clean_key = key[1:] if key.startswith('/') else key
                            metadata[clean_key] = str(value)
                
                # Get basic document properties regardless
                metadata["Pages"] = str(len(reader.pages))
                
                # Fall back to empty metadata rather than error
                if not metadata or list(metadata.keys()) == ["Pages"]:
                    metadata["Info"] = "No metadata found in this PDF file"
                    
                return metadata
        except Exception as e:
            print(f"Warning: Error extracting PDF metadata: {e}")
            # Return minimal info rather than error
            return {"Pages": "Unknown", "Info": "Could not extract metadata"}
    
    def remove_metadata(self, file_path, metadata_type):
        """
        Remove specified metadata from PDF.
        
        Args:
            file_path (str): Path to file
            metadata_type (str): Type of metadata to remove
            
        Returns:
            tuple: (PDF writer object, list of removed tags, None)
        """
        if not PDF_SUPPORT:
            raise Exception("PyPDF2 library not available for PDF metadata removal")
            
        try:
            # Open PDF
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                writer = PdfWriter()
                
                # Copy all pages from the original document
                for i in range(len(reader.pages)):
                    writer.add_page(reader.pages[i])
                
                # For metadata removal, we don't need to do anything else
                # The PdfWriter by default doesn't copy metadata
                
                removed_tags = ["ALL_METADATA"]
                
                return writer, removed_tags, None
            
        except Exception as e:
            print(f"Warning: Error removing PDF metadata: {e}")
            # Try to continue with just copying pages and no metadata
            try:
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    writer = PdfWriter()
                    
                    # Copy all pages from the original document
                    for i in range(len(reader.pages)):
                        writer.add_page(reader.pages[i])
                        
                    return writer, ["ALL_METADATA"], None
            except Exception as e2:
                raise Exception(f"Error processing PDF: {e2}")
    
    def add_metadata(self, file_path, metadata_dict):
        """
        Add or modify metadata in PDF file.
        
        Args:
            file_path (str): Path to file
            metadata_dict (dict): Metadata to add/modify
            
        Returns:
            tuple: (PDF writer object, list of added/modified tags, None)
        """
        if not PDF_SUPPORT:
            raise Exception("PyPDF2 library not available for PDF metadata modification")
            
        try:
            # Open PDF
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                writer = PdfWriter()
                
                # Copy all pages from the original document
                for i in range(len(reader.pages)):
                    writer.add_page(reader.pages[i])
                
                modified_tags = []
                
                # Add metadata to writer - handling the API differences
                if hasattr(writer, 'add_metadata'):
                    writer.add_metadata(metadata_dict)
                else:
                    # For older PyPDF2 versions
                    for k, v in metadata_dict.items():
                        writer.addMetadata(f"/{k}", v)
                
                for key in metadata_dict:
                    modified_tags.append(key)
                
                return writer, modified_tags, None
            
        except Exception as e:
            print(f"Warning: Error adding PDF metadata: {e}")
            # Try to continue with just copying pages and no metadata
            try:
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    writer = PdfWriter()
                    
                    # Copy all pages from the original document
                    for i in range(len(reader.pages)):
                        writer.add_page(reader.pages[i])
                        
                    return writer, [], None
            except Exception as e2:
                raise Exception(f"Error processing PDF: {e2}")
    
    def save_modified_file(self, original_path, writer, _, count):
        """
        Save PDF with modified metadata.
        
        Args:
            original_path (str): Path to original file
            writer: PdfWriter object
            _: Unused (for consistency with interface)
            count (int): File counter for unique filename
            
        Returns:
            str: Path to newly saved file
        """
        try:
            # Create new filename
            base_name = os.path.splitext(original_path)[0]
            new_file = f"{base_name}_modified_{count}.pdf"
            
            # Check if the new file path is too long
            max_length = 250
            if len(new_file) > max_length:
                allowed_base_length = max_length - len(f"_modified_{count}.pdf") - 10
                truncated_base = os.path.splitext(original_path)[0][:allowed_base_length]
                new_file = f"{truncated_base}_modified_{count}.pdf"
            
            # Write to file
            with open(new_file, 'wb') as f:
                writer.write(f)
            
            return new_file
            
        except Exception as e:
            raise Exception(f"Error saving modified PDF: {e}")
    
    def get_metadata_menu_options(self):
        """
        Get menu options specific to PDF files.
        
        Returns:
            list: List of (option_name, metadata_type) tuples
        """
        return [
            ("Remove Author Information", "author"),
            ("Remove Creation Date", "creation_date"),
            ("Remove All Metadata", "all"),
            ("Add/Modify Author", "mod_author"),
            ("Add/Modify Title", "mod_title"),
            ("Add/Modify Subject", "mod_subject"),
            ("Add/Modify Keywords", "mod_keywords"),
            ("Add/Modify Creator", "mod_creator")
        ]


# CLASS Definition
class WordHandler(BaseFileHandler):
    """Handler for Word documents."""
    
    def __init__(self):
        """Initialize WordHandler instance."""
        super().__init__()
        self.file_extension = ['.docx']
        self.display_name = "Word Document"
    
    def validate_file(self, file_path):
        """
        Validate if file is a valid Word document.
        
        Args:
            file_path (str): Path to the file to validate
            
        Returns:
            bool: True if file is valid, False otherwise
            
        Raises:
            ValueError: If the file is corrupted or invalid
        """
        if not super().validate_file(file_path):
            return False
            
        # Check if python-docx is available
        if not DOCX_SUPPORT:
            # Can't validate without the library, assume it's valid
            return True
            
        # Attempt to open the file as Word document
        try:
            doc = Document(file_path)
            # Try to access some properties to verify it's a valid document
            doc.paragraphs
            return True
        except Exception as e:
            error_msg = str(e)
            if any(keyword in error_msg.lower() for keyword in ["corrupt", "invalid", "broken"]):
                print(f"Warning: File validation error: {error_msg}")
                return False
            print(f"Warning: Validation issue, will try to continue: {e}")
            return True
    
    def extract_metadata(self, file_path):
        """
        Extract metadata from Word document with improved error handling.
        
        Args:
            file_path (str): Path to file
            
        Returns:
            dict: Dictionary of metadata
        """
        if not DOCX_SUPPORT:
            return {"Error": "python-docx library not available for Word metadata extraction"}
            
        try:
            doc = Document(file_path)
            metadata = {}
            
            # Extract core properties
            props = doc.core_properties
            
            # Store properties in metadata dictionary
            for prop_name in dir(props):
                # Skip private attributes and methods
                if prop_name.startswith('_') or callable(getattr(props, prop_name)):
                    continue
                
                try:
                    value = getattr(props, prop_name)
                    if value is not None:
                        # Handle datetime objects
                        if isinstance(value, datetime.datetime):
                            metadata[prop_name] = value.strftime("%Y-%m-%d %H:%M:%S")
                        else:
                            metadata[prop_name] = str(value)
                except Exception as e:
                    print(f"Warning: Could not read property {prop_name}: {e}")
            
            return metadata
        except Exception as e:
            print(f"Warning: Error extracting Word metadata: {e}")
            return {}
    
    def remove_metadata(self, file_path, metadata_type):
        """
        Remove specified metadata from Word document.
        
        Args:
            file_path (str): Path to file
            metadata_type (str): Type of metadata to remove
            
        Returns:
            tuple: (Document object, list of removed tags, None)
        """
        if not DOCX_SUPPORT:
            raise Exception("python-docx library not available for Word metadata removal")
            
        try:
            # Open Word document
            doc = Document(file_path)
            removed_tags = []
            
            # Handle different metadata types
            if metadata_type == "all":
                # Get all properties
                props = doc.core_properties
                
                # All writable properties that should be cleared
                writable_props = [
                    'author', 'category', 'comments', 'content_status',
                    'identifier', 'keywords', 'language', 'last_modified_by',
                    'subject', 'title', 'version'
                ]
                
                # Record original metadata for logging
                original_metadata = {}
                for prop_name in dir(props):
                    if not prop_name.startswith('_') and not callable(getattr(props, prop_name)):
                        try:
                            value = getattr(props, prop_name)
                            if value:
                                original_metadata[prop_name] = value
                        except:
                            pass
                
                # Try to clear each property
                for prop_name in writable_props:
                    try:
                        if hasattr(props, prop_name):
                            # Get current value
                            current_value = getattr(props, prop_name)
                            
                            # Only add to removed_tags if there was a value
                            if current_value:
                                removed_tags.append(prop_name)
                                
                            # Set to empty string or None based on type
                            if isinstance(current_value, str):
                                setattr(props, prop_name, "")
                            elif current_value is not None:
                                try:
                                    setattr(props, prop_name, None)
                                except:
                                    setattr(props, prop_name, "")
                    except Exception as e:
                        print(f"Warning: Could not clear property {prop_name}: {e}")
                
                # If no specific tags were identified but we know there was metadata
                if not removed_tags and original_metadata:
                    removed_tags.append("OTHER_METADATA")
                    
                # If we didn't find anything to remove
                if not removed_tags:
                    removed_tags.append("NO_METADATA_FOUND")
            
            elif metadata_type == "author":
                # Remove author information - set to empty string
                try:
                    if doc.core_properties.author:
                        doc.core_properties.author = ""
                        removed_tags.append("author")
                    
                    # Also try to clear last_modified_by if possible
                    if hasattr(doc.core_properties, 'last_modified_by') and doc.core_properties.last_modified_by:
                        doc.core_properties.last_modified_by = ""
                        removed_tags.append("last_modified_by")
                except Exception as e:
                    print(f"Warning: Could not clear author property: {e}")
            
            return doc, removed_tags, None
            
        except Exception as e:
            print(f"Warning: Error during Word metadata removal: {e}")
            try:
                # Try to continue with original document
                doc = Document(file_path)
                return doc, [], None
            except Exception as e2:
                raise Exception(f"Error opening Word document: {e2}")
    
    def add_metadata(self, file_path, metadata_dict):
        """
        Add or modify metadata in Word document.
        
        Args:
            file_path (str): Path to file
            metadata_dict (dict): Metadata to add/modify
            
        Returns:
            tuple: (Document object, list of added/modified tags, None)
        """
        if not DOCX_SUPPORT:
            raise Exception("python-docx library not available for Word metadata modification")
            
        try:
            # Open Word document
            doc = Document(file_path)
            
            modified_tags = []
            
            # Add or modify metadata
            props = doc.core_properties
            for key, value in metadata_dict.items():
                # Skip if property doesn't exist or is read-only
                if not hasattr(props, key) or key in ['created', 'modified']:
                    continue
                
                try:
                    setattr(props, key, value)
                    modified_tags.append(key)
                except Exception as e:
                    print(f"Warning: Could not modify property {key}: {e}")
            
            return doc, modified_tags, None
            
        except Exception as e:
            print(f"Warning: Error during Word metadata addition: {e}")
            try:
                # Try to continue with original document
                doc = Document(file_path)
                return doc, [], None
            except Exception as e2:
                raise Exception(f"Error opening Word document: {e2}")
    
    def save_modified_file(self, original_path, doc, _, count):
        """
        Save Word document with modified metadata.
        
        Args:
            original_path (str): Path to original file
            doc: Document object
            _: Unused (for consistency with interface)
            count (int): File counter for unique filename
            
        Returns:
            str: Path to newly saved file
        """
        try:
            # Create new filename
            base_name = os.path.splitext(original_path)[0]
            new_file = f"{base_name}_modified_{count}.docx"
            
            # Check if the new file path is too long
            max_length = 250
            if len(new_file) > max_length:
                allowed_base_length = max_length - len(f"_modified_{count}.docx") - 10
                truncated_base = os.path.splitext(original_path)[0][:allowed_base_length]
                new_file = f"{truncated_base}_modified_{count}.docx"
            
            # Save document
            doc.save(new_file)
            
            return new_file
            
        except Exception as e:
            raise Exception(f"Error saving modified Word document: {e}")
    
    def get_metadata_menu_options(self):
        """
        Get menu options specific to Word documents.
        
        Returns:
            list: List of (option_name, metadata_type) tuples
        """
        return [
            ("Remove All Metadata", "all"),
            ("Add/Modify Author", "mod_author"),
            ("Add/Modify Title", "mod_title"),
            ("Add/Modify Subject", "mod_subject"),
            ("Add/Modify Keywords", "mod_keywords"),
            ("Add/Modify Category", "mod_category"),
            ("Add/Modify Comments", "mod_comments")
        ]


# CLASS Definition
class ExcelHandler(BaseFileHandler):
    """Handler for Excel spreadsheets."""
    
    def __init__(self):
        """Initialize ExcelHandler instance."""
        super().__init__()
        self.file_extension = ['.xlsx', '.xls']
        self.display_name = "Excel Spreadsheet"
    
    def validate_file(self, file_path):
        """
        Validate if file is a valid Excel spreadsheet.
        
        Args:
            file_path (str): Path to the file to validate
            
        Returns:
            bool: True if file is valid, False otherwise
            
        Raises:
            ValueError: If the file is corrupted or invalid
        """
        if not super().validate_file(file_path):
            return False
            
        # Check if openpyxl is available
        if not EXCEL_SUPPORT:
            # Can't validate without the library, assume it's valid
            return True
            
        # Attempt to open the file as Excel spreadsheet
        try:
            wb = openpyxl.load_workbook(file_path)
            # Try to access the properties to test
            wb.properties
            return True
        except Exception as e:
            error_msg = str(e)
            if any(keyword in error_msg.lower() for keyword in ["corrupt", "invalid", "broken"]):
                print(f"Warning: File validation error: {error_msg}")
                return False
            print(f"Warning: Validation issue, will try to continue: {e}")
            return True
    
    def extract_metadata(self, file_path):
        """
        Extract metadata from Excel spreadsheet with improved error handling.
        
        Args:
            file_path (str): Path to file
            
        Returns:
            dict: Dictionary of metadata
        """
        if not EXCEL_SUPPORT:
            return {"Error": "openpyxl library not available for Excel metadata extraction"}
            
        try:
            wb = openpyxl.load_workbook(file_path, read_only=True)
            metadata = {}
            
            # Extract properties
            props = wb.properties
            
            # Store properties in metadata dictionary
            for prop_name in dir(props):
                # Skip private attributes and methods
                if prop_name.startswith('_') or callable(getattr(props, prop_name)):
                    continue
                
                try:
                    value = getattr(props, prop_name)
                    if value is not None:
                        # Handle datetime objects
                        if isinstance(value, datetime.datetime):
                            metadata[prop_name] = value.strftime("%Y-%m-%d %H:%M:%S")
                        else:
                            metadata[prop_name] = str(value)
                except Exception as e:
                    print(f"Warning: Could not read property {prop_name}: {e}")
            
            return metadata
        except Exception as e:
            print(f"Warning: Error extracting Excel metadata: {e}")
            return {}
    
    def remove_metadata(self, file_path, metadata_type):
        """
        Remove specified metadata from Excel spreadsheet with thorough cleaning.
        
        Args:
            file_path (str): Path to file
            metadata_type (str): Type of metadata to remove
            
        Returns:
            tuple: (Workbook object, list of removed tags, None)
        """
        if not EXCEL_SUPPORT:
            raise Exception("openpyxl library not available for Excel metadata removal")
            
        try:
            # Open Excel spreadsheet
            wb = openpyxl.load_workbook(file_path)
            removed_tags = []
            
            # Get original properties for comparison
            original_props = {}
            if hasattr(wb, 'properties') and wb.properties:
                for prop_name in dir(wb.properties):
                    if not prop_name.startswith('_') and not callable(getattr(wb.properties, prop_name)):
                        try:
                            value = getattr(wb.properties, prop_name)
                            if value is not None:
                                original_props[prop_name] = value
                        except:
                            pass
            
            # Handle different metadata types
            if metadata_type == "all":
                # Create a completely new DocumentProperties object
                new_props = openpyxl.packaging.core.DocumentProperties()
                
                # Record what we're going to remove
                for prop_name, value in original_props.items():
                    # Only record non-system properties that have values
                    if prop_name not in ['namespace', 'tagname', 'idx_base'] and value:
                        removed_tags.append(prop_name)
                
                # Replace the properties with our clean copy
                wb.properties = new_props
                
                # If no specific tags were identified but we know there was metadata
                if not removed_tags and original_props:
                    removed_tags.append("OTHER_METADATA")
                    
                # If nothing was removed
                if not removed_tags:
                    removed_tags.append("NO_METADATA_FOUND")
                    
            elif metadata_type == "author":
                # Remove author information
                props = wb.properties
                
                if hasattr(props, 'creator') and props.creator:
                    props.creator = None
                    removed_tags.append("creator")
                    
                if hasattr(props, 'lastModifiedBy') and props.lastModifiedBy:
                    props.lastModifiedBy = None
                    removed_tags.append("lastModifiedBy")
                
            elif metadata_type == "title":
                # Remove title information
                if hasattr(wb.properties, 'title') and wb.properties.title:
                    wb.properties.title = None
                    removed_tags.append("title")
            
            return wb, removed_tags, None
            
        except Exception as e:
            print(f"Warning: Error during Excel metadata removal: {e}")
            try:
                # Try to continue with original workbook
                wb = openpyxl.load_workbook(file_path)
                return wb, [], None
            except Exception as e2:
                raise Exception(f"Error opening Excel file: {e2}")
    
    def add_metadata(self, file_path, metadata_dict):
        """
        Add or modify metadata in Excel spreadsheet.
        
        Args:
            file_path (str): Path to file
            metadata_dict (dict): Metadata to add/modify
            
        Returns:
            tuple: (Workbook object, list of added/modified tags, None)
        """
        if not EXCEL_SUPPORT:
            raise Exception("openpyxl library not available for Excel metadata modification")
            
        try:
            # Open Excel spreadsheet
            wb = openpyxl.load_workbook(file_path)
            
            modified_tags = []
            
            # Add or modify metadata
            props = wb.properties
            for key, value in metadata_dict.items():
                # Skip if property doesn't exist
                if not hasattr(props, key):
                    continue
                
                try:
                    setattr(props, key, value)
                    modified_tags.append(key)
                except Exception as e:
                    print(f"Warning: Could not modify property {key}: {e}")
            
            return wb, modified_tags, None
            
        except Exception as e:
            print(f"Warning: Error during Excel metadata addition: {e}")
            try:
                # Try to continue with original workbook
                wb = openpyxl.load_workbook(file_path)
                return wb, [], None
            except Exception as e2:
                raise Exception(f"Error opening Excel file: {e2}")
    
    def save_modified_file(self, original_path, wb, _, count):
        """
        Save Excel spreadsheet with modified metadata.
        
        Args:
            original_path (str): Path to original file
            wb: Workbook object
            _: Unused (for consistency with interface)
            count (int): File counter for unique filename
            
        Returns:
            str: Path to newly saved file
        """
        try:
            # Create new filename
            base_name = os.path.splitext(original_path)[0]
            new_file = f"{base_name}_modified_{count}.xlsx"
            
            # Check if the new file path is too long
            max_length = 250
            if len(new_file) > max_length:
                allowed_base_length = max_length - len(f"_modified_{count}.xlsx") - 10
                truncated_base = os.path.splitext(original_path)[0][:allowed_base_length]
                new_file = f"{truncated_base}_modified_{count}.xlsx"
            
            # Save workbook
            wb.save(new_file)
            
            return new_file
            
        except Exception as e:
            raise Exception(f"Error saving modified Excel spreadsheet: {e}")
    
    def get_metadata_menu_options(self):
        """
        Get menu options specific to Excel spreadsheets.
        
        Returns:
            list: List of (option_name, metadata_type) tuples
        """
        return [
            ("Remove Author Information", "author"),
            ("Remove Title", "title"),
            ("Remove All Metadata", "all"),
            ("Add/Modify Creator", "mod_creator"),
            ("Add/Modify Title", "mod_title"),
            ("Add/Modify Subject", "mod_subject"),
            ("Add/Modify Description", "mod_description"),
            ("Add/Modify Keywords", "mod_keywords"),
            ("Add/Modify Category", "mod_category")
        ]